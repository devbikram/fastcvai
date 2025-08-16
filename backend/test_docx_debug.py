#!/usr/bin/env python3
"""
DOCX Debugging Script
Upload your problematic DOCX file to the same directory as this script 
and run: python test_docx_debug.py your_file.docx
"""

import sys
import os
import docx
from docx import Document
import zipfile

def analyze_docx_file(file_path):
    """Comprehensive DOCX file analysis."""
    print(f"🔍 Analyzing DOCX file: {file_path}")
    print("=" * 50)
    
    # Check if file exists
    if not os.path.exists(file_path):
        print(f"❌ File not found: {file_path}")
        return
    
    # Check file size
    file_size = os.path.getsize(file_path)
    print(f"📁 File size: {file_size:,} bytes")
    
    # Check if it's a valid ZIP file (DOCX is a ZIP container)
    try:
        with zipfile.ZipFile(file_path, 'r') as zip_file:
            files_in_docx = zip_file.namelist()
            print(f"📦 Valid ZIP structure with {len(files_in_docx)} files")
            
            # Check for essential DOCX files
            essential_files = ['word/document.xml', '[Content_Types].xml']
            for essential in essential_files:
                if essential in files_in_docx:
                    print(f"✅ Found: {essential}")
                else:
                    print(f"❌ Missing: {essential}")
    except Exception as e:
        print(f"❌ Invalid ZIP structure: {e}")
        return
    
    # Try to open with python-docx
    try:
        doc = Document(file_path)
        print(f"✅ Successfully opened with python-docx")
        
        # Count paragraphs
        paragraph_count = len(doc.paragraphs)
        print(f"📄 Total paragraphs: {paragraph_count}")
        
        # Extract text
        text_parts = []
        non_empty_paragraphs = 0
        
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                text_parts.append(para.text)
                non_empty_paragraphs += 1
                if i < 5:  # Show first 5 non-empty paragraphs
                    print(f"📝 Paragraph {i+1}: '{para.text[:100]}{'...' if len(para.text) > 100 else ''}'")
        
        print(f"📄 Non-empty paragraphs: {non_empty_paragraphs}")
        
        full_text = "\n".join(text_parts)
        text_length = len(full_text)
        
        print(f"📊 Total extracted text length: {text_length} characters")
        
        if text_length == 0:
            print("⚠️  No text extracted! Possible causes:")
            print("   - Document contains only images/tables")
            print("   - Document is password protected")
            print("   - Document has unusual formatting")
            
            # Check for tables
            table_count = len(doc.tables)
            print(f"📋 Tables found: {table_count}")
            
            if table_count > 0:
                print("🔍 Checking table content...")
                for i, table in enumerate(doc.tables[:3]):  # Check first 3 tables
                    for j, row in enumerate(table.rows[:3]):  # Check first 3 rows
                        row_text = " | ".join([cell.text for cell in row.cells])
                        if row_text.strip():
                            print(f"   Table {i+1}, Row {j+1}: {row_text[:100]}")
        else:
            print("✅ Text extraction successful!")
            print(f"📄 Preview (first 200 chars): {full_text[:200]}...")
            
    except Exception as e:
        print(f"❌ Failed to open with python-docx: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python test_docx_debug.py <path_to_docx_file>")
        print("\nThis script will analyze your DOCX file and show detailed information")
        print("to help diagnose why it might not be working with the API.")
        sys.exit(1)
    
    file_path = sys.argv[1]
    analyze_docx_file(file_path)

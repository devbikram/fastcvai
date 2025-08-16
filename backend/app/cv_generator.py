from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from typing import List, Dict, Any

def generate_enhanced_cv(
    original_text: str, 
    additional_info: str, 
    output_path: str,
    improvement_suggestions: Dict[str, Any] = None
) -> None:
    """Generate an enhanced CV in DOCX format with OpenAI-powered improvements."""
    doc = Document()
    
    # Add title with formatting
    title = doc.add_heading('Enhanced Professional CV', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add enhanced professional summary if available
    if improvement_suggestions and "enhanced_summary" in improvement_suggestions:
        doc.add_heading('Professional Summary', level=1)
        summary_para = doc.add_paragraph(improvement_suggestions["enhanced_summary"])
        summary_para.italic = True
        doc.add_paragraph()  # Add space
    
    # Add original CV content with better formatting
    doc.add_heading('Core Experience & Background', level=1)
    
    # Split text into logical sections
    sections = original_text.split('\n\n')
    for section in sections:
        if section.strip():
            lines = section.strip().split('\n')
            if len(lines) > 1 and lines[0].strip():
                # Treat first line as potential section header
                if any(keyword in lines[0].lower() for keyword in ['experience', 'education', 'skills', 'projects', 'certifications']):
                    doc.add_heading(lines[0].strip(), level=2)
                    for line in lines[1:]:
                        if line.strip():
                            doc.add_paragraph(line.strip(), style='List Bullet')
                else:
                    for line in lines:
                        if line.strip():
                            doc.add_paragraph(line.strip())
            else:
                doc.add_paragraph(section.strip())
    
    # Add additional information with enhancements
    if additional_info and additional_info.strip():
        doc.add_heading('Additional Skills & Information', level=1)
        
        # Use improvement suggestions for skill presentation if available
        if improvement_suggestions and "skill_additions" in improvement_suggestions:
            doc.add_paragraph("Enhanced Skill Presentation:", style='Heading 3')
            for suggestion in improvement_suggestions["skill_additions"]:
                doc.add_paragraph(suggestion, style='List Bullet')
        
        doc.add_paragraph("Additional Details:", style='Heading 3')
        additional_sections = additional_info.split('\n\n')
        for section in additional_sections:
            if section.strip():
                doc.add_paragraph(section.strip())
    
    # Add improvement recommendations section
    if improvement_suggestions:
        doc.add_heading('Professional Development Recommendations', level=1)
        
        if "recommendations" in improvement_suggestions:
            doc.add_paragraph("Career Enhancement Suggestions:", style='Heading 3')
            for rec in improvement_suggestions["recommendations"]:
                doc.add_paragraph(rec, style='List Bullet')
        
        if "keyword_optimization" in improvement_suggestions:
            doc.add_paragraph("Industry Keywords to Highlight:", style='Heading 3')
            for keyword in improvement_suggestions["keyword_optimization"]:
                doc.add_paragraph(keyword, style='List Bullet')
        
        if "achievement_focus" in improvement_suggestions:
            doc.add_paragraph("Achievement Enhancement Tips:", style='Heading 3')
            for tip in improvement_suggestions["achievement_focus"]:
                doc.add_paragraph(tip, style='List Bullet')
    
    # Add footer note
    doc.add_paragraph()
    footer_para = doc.add_paragraph("This enhanced CV was generated with AI-powered analysis and recommendations.")
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.runs[0]
    run.italic = True
    run.font.size = Inches(0.1)
    
    # Save the document
    doc.save(output_path)

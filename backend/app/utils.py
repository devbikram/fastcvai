import os
import tempfile
import pdfplumber
import docx
import mammoth
from PIL import Image
from fastapi import UploadFile
from typing import Tuple

# Try to import pytesseract, but make it optional
TESSERACT_AVAILABLE = False
try:
    import pytesseract
    # Test if tesseract is actually working
    pytesseract.get_tesseract_version()
    TESSERACT_AVAILABLE = True
except Exception as e:
    # If any error occurs, Tesseract is not available
    print(f"Tesseract not available: {e}")
    TESSERACT_AVAILABLE = False

def save_temp_file(uploaded_file: UploadFile) -> str:
    """Save uploaded file to temporary location."""
    suffix = os.path.splitext(uploaded_file.filename)[1]
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        content = uploaded_file.file.read()
        tmp.write(content)
        # Reset file pointer for potential re-reading
        uploaded_file.file.seek(0)
        return tmp.name

def extract_text_from_pdf(file_path: str) -> str:
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() or ""
    return text

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX files using mammoth for better formatting support."""
    try:
        # First try mammoth for better text extraction from complex documents
        with open(file_path, "rb") as docx_file:
            result = mammoth.extract_raw_text(docx_file)
            text = result.value.strip()
            
            # If mammoth extracted meaningful text, use it
            if text and len(text) > 10:  # More than just whitespace/minimal content
                return text
                
    except Exception as e:
        print(f"Mammoth extraction failed: {e}, falling back to python-docx")
    
    # Fallback to python-docx method
    try:
        doc = docx.Document(file_path)
        text_parts = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    # Extract text from each cell, including nested paragraphs
                    cell_text = []
                    for para in cell.paragraphs:
                        if para.text.strip():
                            cell_text.append(para.text.strip())
                    if cell_text:
                        row_text.append(" ".join(cell_text))
                if row_text:
                    # Join non-empty cells with " | " and add to text
                    combined_row = " | ".join([text for text in row_text if text])
                    if combined_row.strip():
                        text_parts.append(combined_row)
        
        # If no text found, try extracting from runs (more granular text elements)
        if not text_parts:
            for para in doc.paragraphs:
                for run in para.runs:
                    if run.text.strip():
                        text_parts.append(run.text.strip())
            
            # Also try runs in table cells
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if run.text.strip():
                                    text_parts.append(run.text.strip())
        
        result = "\n".join(text_parts)
        
        # If still no text, this might be an image-based or heavily formatted document
        if not result.strip():
            return "⚠️ No extractable text found. This document may contain only images, complex formatting, or be password protected."
        
        return result
        
    except Exception as e:
        return f"⚠️ Error extracting text from DOCX: {str(e)}"

def extract_text_from_image(file_path: str) -> str:
    """Extract text from image using OCR."""
    # Double-check Tesseract availability at runtime
    try:
        import pytesseract
        pytesseract.get_tesseract_version()
    except Exception as e:
        raise ValueError(
            f"Image text extraction is not available: {str(e)}. "
            "Please install Tesseract OCR: 'brew install tesseract' (macOS) or "
            "visit https://github.com/tesseract-ocr/tesseract for installation instructions."
        )
    
    try:
        image = Image.open(file_path)
        return pytesseract.image_to_string(image)
    except Exception as e:
        raise ValueError(f"Failed to extract text from image: {str(e)}")

def extract_text(file_path: str) -> Tuple[str, str]:
    """Extract text from various file formats."""
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == ".pdf":
        return extract_text_from_pdf(file_path), "pdf"
    elif ext == ".docx":
        return extract_text_from_docx(file_path), "docx"
    elif ext in [".jpg", ".jpeg", ".png"]:
        # Runtime check for Tesseract availability
        try:
            import pytesseract
            pytesseract.get_tesseract_version()
            return extract_text_from_image(file_path), "image"
        except Exception as e:
            raise ValueError(
                f"Image file upload is not supported: {str(e)}. "
                "Please use PDF, DOCX, or TXT files instead. "
                "To enable image support, install Tesseract: 'brew install tesseract' (macOS)"
            )
    elif ext == ".txt":
        return extract_text_from_txt(file_path), "txt"
    else:
        # Check if images would be supported for the error message
        image_support = ""
        try:
            import pytesseract
            pytesseract.get_tesseract_version()
            image_support = " and images (JPG, PNG)"
        except:
            pass
        
        raise ValueError(f"Unsupported file type: {ext}. Supported formats: PDF, DOCX, TXT{image_support}")

def extract_text_from_txt(file_path: str) -> str:
    """Extract text from a plain text file."""
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()
